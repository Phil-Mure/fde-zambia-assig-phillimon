"""
Generate the Sand Healthcare FDE assignment submission as a Word document.
Run: python generate_submission_doc.py

Optional: run capture_screenshots.py first to refresh dashboard images.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from bulletin.services import build_bulletin

PROJECT_ROOT = Path(__file__).resolve().parent
OUTPUT = PROJECT_ROOT / "Sand_FDE_Assignment_Submission.docx"
SCREENSHOT_DIR = PROJECT_ROOT / "submission_assets"
DATA_SOURCE_URL = "https://drive.google.com/drive/folders/1DPk6jKSO_bbnhonUX6S91kLWZVulWTmA"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_image_if_exists(doc: Document, image_path: Path, caption: str, width_inches: float = 6.5) -> None:
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width_inches))
        caption_paragraph = doc.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.runs[0].italic = True
    else:
        doc.add_paragraph(f"[Screenshot pending: {image_path.name}]")


def build_document() -> Document:
    doc = Document()
    sample_report = build_bulletin("2024Q4")

    title = doc.add_heading("Sand Technologies — Forward Deployed Engineer Assignment", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Candidate Submission: Rwanda Health Data Systems Engagement")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_paragraph(
        "This document addresses all four deliverables from the recruitment assignment. "
        "Assumptions are stated explicitly throughout, reflecting the intentional uncertainty "
        "of a Week 1 deployment scenario."
    )

    # ------------------------------------------------------------------ D1
    add_heading(doc, "Deliverable 1: Problem Decomposition and Scoping", 1)

    add_heading(doc, "Section 1: Discovery Process", 2)

    doc.add_paragraph("Who I would talk to in Week 1 and what I would ask:")

    stakeholders = [
        (
            "MoH Director / DG Planning",
            "What decisions does the Quarterly Bulletin inform? Who is the audience? "
            "What is the cost of a wrong or late bulletin?",
        ),
        (
            "DHIS2 system administrator",
            "Which data elements feed the bulletin today? Export process, validation rules, "
            "known data quality issues, reporting timeliness SLAs.",
        ),
        (
            "District Health Officers (2–3 districts)",
            "How do you use facility status information today? What would 'real-time enough' mean "
            "given 3G constraints?",
        ),
        (
            "TB & HIV program leads",
            "Patient matching workflow for co-infection, CommCare export capabilities, "
            "privacy constraints for record linkage.",
        ),
        (
            "HealthTrack EMR site champions (sample hospitals)",
            "Uptime, sync failures, what data never reaches DHIS2.",
        ),
        (
            "Paper facility reporting coordinators",
            "Aggregation path from paper registers to DHIS2, error rates, lag.",
        ),
        (
            "Solutions Manager & Country Director",
            "Political priorities, MoU success criteria, budget/timeline for 6-week sprint.",
        ),
    ]
    for role, questions in stakeholders:
        p = doc.add_paragraph()
        p.add_run(f"{role}: ").bold = True
        p.add_run(questions)

    doc.add_paragraph("Patterns I would look for when decomposing 'our data is a mess':")
    add_bullets(
        doc,
        [
            "Duplication — same indicator reported differently across DHIS2, EMR, Excel.",
            "Timeliness gaps — bulletin uses Q-1 data compiled in week 4 of the new quarter.",
            "Completeness — rural paper facilities missing from national aggregates.",
            "Lineage blindness — no documented path from source register to bulletin table.",
            "Decision latency — officers acting on 3-week-old DHIS2 vs operational reality.",
            "Human bottleneck — 40 hrs/month manual Excel work = single point of failure.",
            "Fragmented programs — TB/HIV silos with no shared patient identifier strategy.",
        ],
    )

    doc.add_paragraph("What I would observe directly (not just hear about):")
    add_bullets(
        doc,
        [
            "Sit in on one manual bulletin compilation session — count copy-paste steps.",
            "Pull a live DHIS2 export and compare to published bulletin figures.",
            "Visit one hospital (HealthTrack) and one rural paper facility same week.",
            "Test connectivity and DHIS2 load time from a district office on 3G.",
            "Review CommCare export format and whether patient IDs are linkable.",
            "Check server uptime logs at 2–3 HealthTrack sites if accessible.",
        ],
    )

    add_heading(doc, "Section 2: Problem Selection", 2)

    doc.add_paragraph(
        "Selected problem for 6-week sprint: Problem A — Automate the Quarterly Health Bulletin"
    ).runs[0].bold = True

    doc.add_paragraph("Why Problem A over B and C:")
    add_bullets(
        doc,
        [
            "Problem A — Highest near-term ROI: 40 hrs/month MoH staff time recovered; "
            "aligns with signed MoU language on 'decision-making capabilities'; "
            "uses DHIS2 data already aggregated nationally; achievable in 6 weeks with "
            "Sand Analytics Template Toolkit + custom ETL.",
            "Problem B — Higher infrastructure risk (175 paper facilities, spotty connectivity, "
            "buggy EMRs). Better as Phase 2 once data pipeline trust exists. Health Atlas is the "
            "right product, but real-time facility status needs reliable ingestion first.",
            "Problem C — Critical clinically, but requires patient identity resolution, "
            "governance, and TB/HIV program buy-in. 6 weeks is insufficient for ethical, "
            "compliant record linkage across CommCare instances.",
        ],
    )

    doc.add_paragraph("Specific, measurable outcome (6-week sprint):")
    add_bullets(
        doc,
        [
            "Automated bulletin covering Q1 metrics generated in < 30 minutes (vs 40 hrs manual).",
            "HTML dashboard + exportable JSON/PDF consumed by MoH Planning unit.",
            "Top 10 facilities, maternal indicators, performance scores, and QoQ trends — "
            "matching manual bulletin within 2% variance on pilot validation.",
        ],
    )

    doc.add_paragraph("Explicitly OUT OF SCOPE (Sprint 1):")
    add_bullets(
        doc,
        [
            "Real-time facility status / stockout mapping (Problem B).",
            "Cross-program patient matching for TB/HIV co-infection (Problem C).",
            "Full HealthTrack / OpenMRS integration — DHIS2 monthly aggregates only.",
            "175 rural paper facility digitization.",
            "Production HA deployment, SSO, or full Superset self-service rollout.",
        ],
    )

    doc.add_paragraph("Key assumptions and Week 2 validation:")
    add_bullets(
        doc,
        [
            "Assumption: The five provided source files contain enough structure to model "
            "a quarterly neonatal bulletin while DHIS2 API access is pending.",
            "Validate: Compare prototype aggregates to the last published bulletin with MoH analysts.",
            "Assumption: Composite readiness scoring weights (35% outcomes / 20% governance / "
            "20% workforce / 15% operations / 10% infrastructure) are directionally acceptable.",
            "Validate: Workshop with Planning unit to tune weights and thresholds.",
            "Assumption: Monthly clinical rows can be rolled into calendar quarters for bulletin periods.",
            "Validate: Confirm quarter definitions and partial-quarter handling with DHIS2 administrators.",
        ],
    )

    doc.add_paragraph("Success metric:")
    doc.add_paragraph(
        "MoH Planning lead signs off that automated bulletin replaces manual compilation for "
        "one full quarter, with ≤ 2% numeric variance and generation time < 30 minutes."
    )

    doc.add_paragraph("Fallback plan if harder than expected:")
    add_bullets(
        doc,
        [
            "Week 3 checkpoint: If DHIS2 API access is delayed, continue with scheduled CSV "
            "imports (current prototype approach) and document API integration as Week 4–6 goal.",
            "If performance scoring is contested, deliver raw timeliness/completeness tables first, "
            "defer composite score to MoH-configurable rules.",
            "If PDF layout is blocked, ship HTML dashboard + JSON to Superset; manual PDF paste "
            "still saves > 30 hrs/month.",
        ],
    )

    # ------------------------------------------------------------------ D2
    add_heading(doc, "Deliverable 2: Rapid Prototyping (Problem A)", 1)

    add_heading(doc, "Section 1: Solution Design & Architecture", 2)

    doc.add_paragraph("Architecture overview (data flow):")
    doc.add_paragraph(
        "[Google Drive Source Files: clinical_neonatal | facilities | governance | "
        "healthcare_workers | operations] → [Drive Sync + CSV/XLSX Loader] → "
        "[HealthOS-style Join Layer] → [Bulletin Metrics Engine] → "
        "[Outputs: Executive HTML Dashboard | JSON API | Export Payload]"
    ).runs[0].font.name = "Consolas"

    doc.add_paragraph("Source data folder (assignment dataset):")
    doc.add_paragraph(DATA_SOURCE_URL)

    doc.add_paragraph("Sand products leveraged:")
    add_bullets(
        doc,
        [
            "HealthOS Data Models — normalize five facility-level source files into one "
            "joined reporting schema before aggregation.",
            "Analytics Template Toolkit (Apache Superset) — Phase 2 self-service exploration; "
            "prototype exposes JSON API for Superset dataset connection.",
            "Health Insight Engine — future: anomaly flags on neonatal mortality spikes and "
            "stockout clusters (risk watchlist is the Week 2 precursor).",
            "Health Atlas — future: map provincial readiness and referral burden once Problem B starts.",
        ],
    )

    doc.add_paragraph("Custom build (justification):")
    add_bullets(
        doc,
        [
            "Django bulletin service — MoH-specific metric definitions, quarter logic, and bulletin layout.",
            "Google Drive ingestion layer — keeps the prototype reproducible from the exact assignment dataset.",
            "Multi-source join and composite scoring — combines clinical, governance, workforce, and operations signals.",
            "Executive dashboard UI — ministry-facing narrative view not available off-the-shelf in Superset alone.",
        ],
    )

    doc.add_paragraph("Build vs buy decisions:")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Decision"
    hdr[2].text = "Rationale"
    rows = [
        ("Data warehouse", "Buy — HealthOS", "Reuse Sand transformations"),
        ("BI dashboards", "Buy — Superset templates", "MoH analysts already need self-service"),
        ("Bulletin orchestration", "Build — Django", "Specific MoH bulletin format"),
        ("DHIS2 connector", "Buy/adapt — existing HealthOS connector", "Don't rebuild API client"),
        ("Auth / SSO", "Buy — MoH IdP integration", "Defer to hardening phase"),
        ("PDF generation", "Build light — WeasyPrint/ReportLab", "Or export to MoH Word template"),
    ]
    for comp, decision, rationale in rows:
        row = table.add_row().cells
        row[0].text = comp
        row[1].text = decision
        row[2].text = rationale

    add_heading(doc, "Section 2: Working Prototype CODE", 2)

    doc.add_paragraph(
        "Working Django project location: sand-fde-assignment/ (submitted with this document)"
    )
    doc.add_paragraph("Prototype capabilities:")
    add_bullets(
        doc,
        [
            "Reads only the five provided Google Drive files (CSV or Excel) from data/google_drive_source/",
            "Auto-downloads the shared folder on first run if files are not already present locally",
            "Calculates bulletin metrics: top facilities, maternal/neonatal indicators, composite performance scores, provincial overview, monthly trends, QoQ momentum, and risk watchlist",
            "Executive HTML dashboard at /",
            "JSON API at /api/bulletin/",
            "Downloadable JSON export at /export/",
        ],
    )

    doc.add_paragraph("Sample output from the current prototype (2024Q4):")
    add_bullets(
        doc,
        [
            f"Facilities covered: {sample_report.source_summary['facility_count']}",
            f"Quarter deliveries: {sample_report.maternal_summary['deliveries']:,}",
            f"Neonatal mortality rate: {sample_report.maternal_summary['neonatal_mortality_rate']}/1k live births",
            f"Top facility by volume: {sample_report.top_facilities[0]['facility_name']} ({sample_report.top_facilities[0]['deliveries']:,} deliveries)",
            f"Highest composite readiness score: {sample_report.performance_scores[0]['facility_name']} ({sample_report.performance_scores[0]['overall_score']})",
            f"Provinces benchmarked: {len(sample_report.provincial_overview)}",
            f"Facilities flagged on risk watchlist: {len(sample_report.risk_alerts)}",
        ],
    )

    doc.add_paragraph("Setup and run:")
    add_numbered(
        doc,
        [
            "cd sand-fde-assignment",
            "python -m venv .venv",
            ".venv\\Scripts\\activate  (Windows) or source .venv/bin/activate (macOS/Linux)",
            "pip install -r requirements.txt",
            "python manage.py migrate",
            "python manage.py test",
            "python manage.py runserver",
            "Open http://127.0.0.1:8000/",
        ],
    )

    doc.add_paragraph(
        "Key code modules: bulletin/services.py (Drive sync, joins, scoring, metrics), "
        "bulletin/views.py (dashboard/API), bulletin/templates/bulletin/dashboard.html "
        "(executive MoH-facing output), capture_screenshots.py (submission screenshots)."
    )

    add_heading(doc, "Dashboard screenshots", 3)
    doc.add_paragraph(
        "The screenshots below are generated from the live prototype so reviewers can see the "
        "actual working output without running the server."
    )
    add_image_if_exists(
        doc,
        SCREENSHOT_DIR / "dashboard_hero.png",
        "Figure 1 — Executive dashboard hero, KPI strip, and strategic summary cards",
    )
    add_image_if_exists(
        doc,
        SCREENSHOT_DIR / "dashboard_tables.png",
        "Figure 2 — Facility rankings, risk watchlist, provincial overview, and performance leaders",
    )

    add_heading(doc, "Section 3: Implementation Notes", 2)

    doc.add_paragraph("Shortcuts taken to move fast:")
    add_bullets(
        doc,
        [
            "Google Drive folder sync instead of live DHIS2 API — avoids credential blockers in assignment timeframe.",
            "Composite readiness score uses transparent, code-defined weights — not yet MoH-configurable in UI.",
            "No database persistence — stateless compute from files; production would use Postgres or HealthOS models.",
            "JSON export instead of formatted PDF — proves data pipeline; PDF is template work for Week 4–6.",
            "Latest quarter may be partial when the shared clinical file does not yet contain all three months.",
        ],
    )

    doc.add_paragraph("Week 3 demo — what works vs what is broken:")
    add_bullets(
        doc,
        [
            "Works: Multi-source ingestion, quarter selector, KPI cards, maternal/neonatal indicators, top facilities, performance leaders, provincial overview, risk watchlist, JSON API/export, unit tests.",
            "Broken / stubbed: Live DHIS2 pull, user authentication, scheduled monthly refresh, email distribution to districts, Superset dataset auto-refresh.",
            "Fragile: Composite score weights are hard-coded; changing MoH policy requires a code edit or config layer.",
        ],
    )

    doc.add_paragraph("What I learned from building:")
    add_bullets(
        doc,
        [
            "A credible bulletin needs more than one file — joining clinical, governance, workforce, and operations data creates actionable readiness signals.",
            "Separating ingestion, scoring, and presentation enables HTML, JSON, and Superset reuse from one metrics engine.",
            "Risk watchlists surface intervention priorities faster than tables alone — mortality plus stockouts plus staffing gaps tell a clearer story.",
            "Django + Pandas is sufficient for rapid FDE prototypes; migration path to SHOS services and scheduled ETL is clear.",
        ],
    )

    # ------------------------------------------------------------------ D3
    add_heading(doc, "Deliverable 3: Production Hardening", 1)

    hardening = [
        (
            "DHIS2 integration & data freshness",
            "Concern: CSV manual drops reintroduce the bottleneck and stale data risk.",
            "Address: Scheduled DHIS2 API pull via HealthOS connector; incremental sync; "
            "data freshness indicator on dashboard; alert if export > 21 days old.",
        ),
        (
            "Security & access control",
            "Concern: Health data exposed without role-based access.",
            "Address: MoH SSO integration; district-level row filters; audit logging; "
            "HTTPS termination; secrets in vault not repo.",
        ),
        (
            "Data quality & validation",
            "Concern: Composite readiness scores depend on proxy fields and partial quarters.",
            "Address: Implement DHIS2 validation rules; anomaly detection via Health Insight Engine; "
            "manual override workflow for district reviewers; explicit partial-quarter badges.",
        ),
        (
            "Reliability & observability",
            "Concern: Silent pipeline failures during unreliable power/connectivity.",
            "Address: Idempotent ETL jobs; retry queues; dead-letter exports; "
            "monitoring dashboards; on-call runbook for MoH IT.",
        ),
        (
            "Output fidelity & distribution",
            "Concern: JSON export is not the official bulletin format MoH publishes.",
            "Address: PDF generation matching MoH template; automated email to district lists; "
            "versioned bulletin archive; Superset dashboards for drill-down.",
        ),
    ]
    for i, (title, concern, address) in enumerate(hardening, 1):
        add_heading(doc, f"{i}. {title}", 2)
        doc.add_paragraph(f"Concern: {concern}")
        doc.add_paragraph(f"Address: {address}")

    # ------------------------------------------------------------------ D4
    add_heading(doc, "Deliverable 4: Handover & Deployment", 1)

    doc.add_paragraph("Three things to document for MoH IT:")
    add_numbered(
        doc,
        [
            "Data pipeline runbook — DHIS2 credentials, cron schedule, failure recovery, "
            "contact escalation (Sand + MoH).",
            "Metric definitions document — each bulletin field with formula, source data element, "
            "and known limitations.",
            "Environment & deployment guide — server requirements, Django/Superset versions, "
            "backup strategy, and how to regenerate a bulletin manually.",
        ],
    )

    doc.add_paragraph("Training plan:")
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    h2 = table2.rows[0].cells
    h2[0].text = "Audience"
    h2[1].text = "Topics"
    h2[2].text = "Format"
    training = [
        ("MoH Planning analysts", "Running bulletin, interpreting trends, exporting PDF", "2 hr workshop"),
        ("DHIS2 administrators", "Data quality checks, validation rules, API tokens", "Half-day technical"),
        ("District health officers", "Reading dashboard, identifying at-risk facilities", "1 hr webinar per zone"),
        ("MoH IT / infrastructure", "Deployment, monitoring, incident response", "Hands-on with Sand FDE"),
    ]
    for aud, topics, fmt in training:
        r = table2.add_row().cells
        r[0].text = aud
        r[1].text = topics
        r[2].text = fmt

    doc.add_paragraph("Exit criteria (transition out):")
    add_bullets(
        doc,
        [
            "Two consecutive quarters generated by MoH staff without Sand FDE on-call support.",
            "All runbooks and metric definitions signed off by MoH IT and Planning leads.",
            "Superset dashboards adopted by at least 3 district officers with positive feedback.",
            "Incident SLA agreed: MoH IT owns L1; Sand central FDE available for L2 escalations only.",
            "Backlog for Problems B (Health Atlas) and C (program linkage) handed to Solutions Manager.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "Appendix: Submission Checklist", 1)
    add_bullets(
        doc,
        [
            "Scoping and design documents — this Word file",
            "Working code with provided Google Drive sample data — sand-fde-assignment/",
            "README with setup/run instructions — README.md",
            "Output demonstration — run server and open http://127.0.0.1:8000/ or review embedded screenshots",
            "Screenshot assets — submission_assets/ (generated via capture_screenshots.py)",
        ],
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
