"""
Generate the final Sand Healthcare FDE assignment submission for Phillimon Murebwa.

The document is intentionally generated from code so the submission is reproducible
and stays aligned with the working prototype metrics.
"""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from bulletin.services import available_periods, build_bulletin

PROJECT_ROOT = Path(__file__).resolve().parent
OUTPUT_MAIN = PROJECT_ROOT / "Sand_Healthcare_Forward_Deployed_Engineer_Assignment_1125-2.docx"
OUTPUT_COMPAT = PROJECT_ROOT / "Sand_FDE_Assignment_Submission.docx"
OUTPUT_DOWNLOADS = PROJECT_ROOT.parent / "Sand_Healthcare_Forward_Deployed_Engineer_Assignment_1125-2.docx"
SCREENSHOT_DIR = PROJECT_ROOT / "submission_assets"
DATA_SOURCE_URL = "https://drive.google.com/drive/folders/1DPk6jKSO_bbnhonUX6S91kLWZVulWTmA"

NAVY = "0F172A"
TEAL = "0B6E8E"
GREEN = "0D9488"
LIGHT_TEAL = "E6F4F8"
MINT = "CCFBF1"
SLATE = "334155"
GRAY = "F8FAFC"
WHITE = "FFFFFF"
AMBER = "FFFAEB"
RED = "FEF3F2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D7E0EA", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: str = NAVY, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top: int = 100, start: int = 100, bottom: int = 100, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 25, NAVY),
        ("Heading 1", 16, TEAL),
        ("Heading 2", 12, NAVY),
        ("Heading 3", 10, SLATE),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    styles["List Bullet"].font.name = "Aptos"
    styles["List Number"].font.name = "Aptos"


def add_title(doc: Document, text: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SAND HEALTHCARE FORWARD DEPLOYED ENGINEER")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(text)
    title_run.bold = True
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor.from_string(NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(subtitle)
    sub_run.font.name = "Aptos"
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor.from_string(SLATE)


def add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "B8C5D6")
    set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run("\n")
    b = p.add_run(body)
    b.font.name = "Aptos"
    b.font.size = Pt(9)
    b.font.color.rgb = RGBColor.from_string(NAVY)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], title: str | None = None) -> None:
    if title:
        add_section_heading(doc, title, 3)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["Area", "Answer"]):
        set_cell_shading(cell, TEAL)
        set_cell_text(cell, text, bold=True, color=WHITE, size=9)
        set_cell_border(cell, TEAL)
    for idx, (left, right) in enumerate(rows):
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT_TEAL if idx % 2 == 0 else GRAY)
        set_cell_shading(cells[1], WHITE if idx % 2 == 0 else GRAY)
        for cell in cells:
            set_cell_border(cell)
            set_cell_margins(cell, 90, 90, 90, 90)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor.from_string(NAVY)
        cells[0].paragraphs[0].runs[0].bold = True


def add_metrics_strip(doc: Document, metrics: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, value, note) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, [LIGHT_TEAL, MINT, GRAY, AMBER][idx % 4])
        set_cell_border(cell, "B8C5D6")
        set_cell_margins(cell, 140, 110, 140, 110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label.upper())
        r1.bold = True
        r1.font.name = "Aptos"
        r1.font.size = Pt(7)
        r1.font.color.rgb = RGBColor.from_string(SLATE)
        p.add_run("\n")
        r2 = p.add_run(value)
        r2.bold = True
        r2.font.name = "Aptos Display"
        r2.font.size = Pt(15)
        r2.font.color.rgb = RGBColor.from_string(TEAL)
        p.add_run("\n")
        r3 = p.add_run(note)
        r3.font.name = "Aptos"
        r3.font.size = Pt(7)
        r3.font.color.rgb = RGBColor.from_string(SLATE)


def add_architecture_diagram(doc: Document) -> None:
    add_section_heading(doc, "Architecture overview: data flow from messy source systems to decisions", 3)
    add_callout(
        doc,
        "Design principle",
        "I modelled the flow from the lowest reliable reporting unit upward: monthly DHIS2-like rows are "
        "validated and joined first, then rolled into quarterly and annual bulletin views. Counts are summed "
        "before rates are recalculated, so annual neonatal mortality is not an average of quarterly rates.",
        fill=GRAY,
    )

    headers = [
        "1. Source systems",
        "2. Ingestion layer",
        "3. HealthOS data models",
        "4. Bulletin metrics engine",
        "5. Outputs and actions",
    ]
    body = [
        "DHIS2 monthly exports\nAssignment CSV/XLSX files\nHealthTrack EMR and OpenMRS future feeds\nPaper facility registers via district upload\nTB/HIV/Immunisation feeds for later phases",
        "Google Drive sync with local cache\nCSV/XLSX loader\nSchema and role classification\nPartial-period detection\nFuture DHIS2 API connector with retries",
        "Facility dimension\nReporting period dimension\nClinical neonatal fact table\nGovernance, workforce, operations, infrastructure dimensions\nSupplemental facility-level joins",
        "Monthly -> quarterly -> annual filtering\nTop facilities, maternal/neonatal indicators\nComposite readiness scores\nMoM/QoQ/YoY trend analysis\nRisk watchlist and rollup reconciliation",
        "Executive HTML dashboard\nJSON API and export payload\nSuperset dataset feed\nMoH PDF/Word bulletin template\nFuture Health Atlas maps and Insight Engine alerts",
    ]
    products = [
        "External systems",
        "Custom sprint build",
        "HealthOS Data Models",
        "Custom + Health Insight pattern",
        "Analytics Toolkit + Health Atlas path",
    ]

    table = doc.add_table(rows=3, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    palette = [LIGHT_TEAL, MINT, "E0F2FE", "EEF2FF", AMBER]

    for idx, header in enumerate(headers):
        col = idx * 2
        cell = table.cell(0, col)
        set_cell_shading(cell, TEAL if idx in {1, 3} else NAVY)
        set_cell_text(cell, header, bold=True, color=WHITE, size=8)
        set_cell_border(cell, TEAL if idx in {1, 3} else NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if col + 1 < 9:
            arrow = table.cell(0, col + 1)
            set_cell_text(arrow, "->", bold=True, color=TEAL, size=12)
            set_cell_border(arrow, WHITE)

        body_cell = table.cell(1, col)
        set_cell_shading(body_cell, palette[idx])
        set_cell_text(body_cell, body[idx], color=NAVY, size=7)
        set_cell_border(body_cell, "B8C5D6")
        set_cell_margins(body_cell, 120, 80, 120, 80)
        body_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if col + 1 < 9:
            arrow = table.cell(1, col + 1)
            set_cell_text(arrow, "->", bold=True, color=TEAL, size=12)
            set_cell_border(arrow, WHITE)

        product_cell = table.cell(2, col)
        set_cell_shading(product_cell, GRAY)
        set_cell_text(product_cell, products[idx], bold=True, color=SLATE, size=7)
        set_cell_border(product_cell, "D7E0EA")
        if col + 1 < 9:
            spacer = table.cell(2, col + 1)
            set_cell_border(spacer, WHITE)

    doc.add_paragraph(
        "Current prototype path: Google Drive source files -> Drive sync + pandas loader -> HealthOS-style joined "
        "facility model -> Django metrics engine -> HTML dashboard, JSON API, and export payload."
    )
    doc.add_paragraph(
        "Production path: DHIS2 API and HealthTrack/OpenMRS connectors feed HealthOS models on a schedule; "
        "the same metrics definitions power Superset templates, official PDF bulletins, Health Atlas maps, and "
        "Health Insight Engine alerts."
    )


def add_image_if_exists(doc: Document, image_path: Path, caption: str, width_inches: float = 6.7) -> None:
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=Inches(width_inches))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_p.add_run(caption)
    run.italic = True
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SLATE)


def build_document() -> Document:
    latest_month = available_periods(granularity="M")[-1].key
    monthly_report = build_bulletin(latest_month, granularity="M")
    quarterly_report = build_bulletin("2024Q4", granularity="Q")
    annual_report = build_bulletin("2024", granularity="A")

    doc = Document()
    style_document(doc)
    doc.core_properties.author = "Phillimon Murebwa"
    doc.core_properties.title = "Sand Healthcare Forward Deployed Engineer Assignment"
    doc.core_properties.subject = "Rwanda Health Data Systems - Problem A Quarterly Health Bulletin Automation"

    add_title(
        doc,
        "Rwanda Health Data Systems Engagement",
        "Forward Deployed Engineer assignment submission by Phillimon Murebwa",
    )
    add_metrics_strip(
        doc,
        [
            ("Selected sprint", "Problem A", "Quarterly Health Bulletin automation"),
            ("Prototype status", "Working", "Django dashboard + JSON API"),
            ("Data sources", "5 files", "Clinical, facilities, governance, workforce, operations"),
            ("Reporting units", "M/Q/A", "Monthly, quarterly, annual rollups"),
        ],
    )
    add_callout(
        doc,
        "Executive summary",
        "I chose Problem A because it is the fastest route to measurable MoH value: it targets a known 40-hour/month "
        "manual bottleneck, uses already-available DHIS2-like aggregate data, and creates the trust foundation needed "
        "before moving into real-time facility status or cross-program patient matching. I built a working Django "
        "prototype that ingests the provided source files, joins them into a HealthOS-style model, calculates bulletin "
        "metrics, supports advanced monthly/quarterly/annual filtering, and exposes both an executive dashboard and JSON API.",
        fill=LIGHT_TEAL,
    )
    doc.add_paragraph(
        f"Generated on {datetime.now():%d %B %Y}. Source dataset used: {DATA_SOURCE_URL}"
    )
    doc.add_page_break()

    add_section_heading(doc, "Assignment Compliance Checklist", 1)
    checklist = [
        ("Deliverable 1 - Discovery process", "Complete: stakeholders, questions, patterns, and direct observations are listed."),
        ("Deliverable 1 - Problem selection", "Complete: Problem A selected with measurable outcome, scope, assumptions, success metric, and fallback."),
        ("Deliverable 2 - Architecture", "Complete: includes a visual data-flow diagram and build-vs-buy table."),
        ("Deliverable 2 - Working code", "Complete: Django + pandas prototype reads provided data, calculates metrics, and outputs HTML/JSON."),
        ("Deliverable 2 - Implementation notes", "Complete: shortcuts, Week 3 demo plan, and lessons learned."),
        ("Deliverable 3 - Production hardening", "Complete: top 5 production fixes with concern and remediation."),
        ("Deliverable 4 - Handover", "Complete: MoH IT documentation, training plan, and exit criteria."),
        ("Submission evidence", "Complete: README, tests, screenshots, source files, and regenerated Word submission."),
    ]
    add_key_value_table(doc, checklist)

    add_section_heading(doc, "Deliverable 1: Problem Decomposition and Scoping", 1)
    add_section_heading(doc, "Section 1: Discovery Process", 2)
    doc.add_paragraph(
        "My Week 1 goal would be to translate 'our data is a mess' into a ranked set of decision failures, data-flow "
        "breakpoints, and measurable delivery opportunities. I would avoid starting with tooling. I would start with "
        "people, decisions, source systems, and the manual work currently keeping the system alive."
    )
    add_key_value_table(
        doc,
        [
            ("MoH Director / DG Planning", "Which decisions does the bulletin actually drive? What decisions are delayed or wrong today? What would convince you in six weeks that Sand created value?"),
            ("MoH Planning analysts", "Walk me through the current 40-hour compilation process. Which Excel steps are repetitive, fragile, or politically sensitive?"),
            ("DHIS2 administrator", "Which data elements feed the bulletin, what are the validation rules, and where do late or incomplete reports appear?"),
            ("District Health Officers", "What do you do when a facility appears high-risk? What freshness is good enough for district action under 3G constraints?"),
            ("Hospital and clinic data clerks", "Where do HealthTrack/OpenMRS values diverge from DHIS2 submissions, and what manual reconciliation happens locally?"),
            ("Paper facility reporting coordinators", "How do registers become district uploads, what gets lost, and how long is the lag?"),
            ("TB/HIV/Immunisation leads", "What identifiers exist across CommCare/Excel systems, and what governance is required before any patient-level linkage?"),
            ("Sand country team", "What is politically urgent, what is reusable across countries, and where should I ask for central FDE/Product support?"),
        ],
        title="Stakeholders I would interview and the questions I would ask",
    )
    add_section_heading(doc, "Patterns I would look for", 3)
    add_bullets(
        doc,
        [
            "Decision latency: reports arrive weeks after action windows have passed.",
            "Manual transformation risk: copy-paste Excel steps that are undocumented or owned by one analyst.",
            "Metric drift: different teams define the same indicator differently across DHIS2, EMR, Excel, and paper workflows.",
            "Completeness gaps: rural paper facilities and late monthly reports create invisible denominator problems.",
            "Lineage blindness: leaders see final numbers but cannot trace them back to source facility, month, or data element.",
            "Program silos: maternal/neonatal, TB, HIV, immunisation, and operations teams optimize separate datasets.",
            "Infrastructure constraints: unreliable power and spotty 3G/4G change what 'real-time' can realistically mean.",
        ],
    )
    add_section_heading(doc, "What I would observe directly", 3)
    add_bullets(
        doc,
        [
            "Sit beside the analyst compiling the bulletin and count every manual export, pivot, formula, paste, and validation step.",
            "Compare one DHIS2 export to the last published bulletin and identify variance by metric and facility.",
            "Visit one HealthTrack hospital, one OpenMRS clinic, and one paper-only rural facility in the same reporting cycle.",
            "Measure actual district-office connectivity and dashboard load time under normal 3G/4G conditions.",
            "Inspect a CommCare export from TB/HIV programs to understand whether future linkage is feasible and ethical.",
            "Review late-report and missing-report logs to distinguish data-quality issues from true service-performance issues.",
        ],
    )

    add_section_heading(doc, "Section 2: Problem Selection", 2)
    add_callout(
        doc,
        "Selected 6-week sprint: Problem A - automate the Quarterly Health Bulletin",
        "I would choose Problem A because it is bounded, visible to senior MoH stakeholders, measurable in hours saved, "
        "and technically achievable with aggregate data in six weeks. It also creates a reusable data foundation for "
        "Problem B (Health Atlas facility status) and Problem C (program linkage) without prematurely taking on their "
        "higher infrastructure, privacy, and workflow risks.",
        fill=MINT,
    )
    add_key_value_table(
        doc,
        [
            ("Why not Problem B first?", "Real-time facility status is valuable, but 175 paper-only facilities, unreliable power, and spotty 3G/4G make it a broader operational change. I would treat it as Phase 2 after proving data trust."),
            ("Why not Problem C first?", "TB/HIV co-infection visibility is clinically important, but patient matching across CommCare systems requires identity governance, privacy controls, and program buy-in. Six weeks is too short to do this safely."),
            ("Specific outcome", "Generate an automated bulletin for one full reporting cycle in under 30 minutes, replacing the 40-hour manual workflow and matching validated MoH figures within an agreed tolerance."),
            ("Out of scope", "Patient-level record linkage, full rural paper digitisation, production HA deployment, SSO, national Health Atlas rollout, and official PDF distribution automation."),
            ("Week 2 validations", "Confirm indicator definitions, quarter boundaries, partial-period handling, acceptable variance, sign-off workflow, and whether DHIS2 API access is available or CSV export remains the first ingestion path."),
            ("Success metric", "MoH Planning lead accepts an automated bulletin for a pilot quarter with generation time below 30 minutes and traceable metric definitions."),
            ("Fallback plan", "If DHIS2 API access is delayed, continue with scheduled CSV/Excel ingestion; if composite scoring is debated, ship raw indicator tables plus configurable weights; if PDF is delayed, ship HTML and JSON first."),
        ],
    )

    doc.add_page_break()
    add_section_heading(doc, "Deliverable 2: Rapid Prototyping - Problem A", 1)
    add_section_heading(doc, "Section 1: Solution Design and Architecture", 2)
    add_architecture_diagram(doc)

    add_section_heading(doc, "Sand products leveraged", 3)
    add_key_value_table(
        doc,
        [
            ("HealthOS Data Models", "Normalize facility, period, clinical, governance, workforce, operations, and infrastructure facts before any dashboard calculation."),
            ("Analytics Template Toolkit", "Use Apache Superset templates for Phase 2 analyst self-service while the prototype provides immediate HTML/JSON outputs."),
            ("Health Insight Engine", "Production path for anomaly detection, mortality spikes, stockout clusters, and automated district alerts."),
            ("Health Atlas", "Phase 2 map layer for facility status, provincial readiness, stockouts, outbreak signals, and referral burden."),
            ("Health Outcome Tracker", "Future patient-outcome analytics once Problem C identity governance is safe enough to pursue."),
        ],
    )
    add_section_heading(doc, "Custom build and justification", 3)
    add_bullets(
        doc,
        [
            "Django service layer for MoH-specific bulletin logic, period selection, JSON/export routes, and dashboard rendering.",
            "Pandas metrics engine for rapid multi-source joins, rollups, rate calculations, composite scoring, and risk ranking.",
            "Google Drive/CSV/XLSX ingestion layer so the assignment remains reproducible without DHIS2 credentials.",
            "Executive dashboard UI because the MoH needs a narrative action view, not only a generic BI grid.",
            "Advanced filtering logic from monthly source rows to quarterly and annual views, because the real DHIS2 workflow starts with monthly reporting.",
        ],
    )
    add_key_value_table(
        doc,
        [
            ("Data warehouse/model", "Buy/adapt HealthOS Data Models. Reuse Sand healthcare transformations instead of inventing a national model from scratch."),
            ("Ingestion", "Build thin adapter now; buy/adapt DHIS2 connector in production. The sprint needs speed, production needs reliability."),
            ("Metric engine", "Build custom. MoH indicator definitions, partial periods, and bulletin layout are engagement-specific."),
            ("Dashboard", "Build focused executive view now; feed Superset later for self-service."),
            ("Maps and alerts", "Buy/use Health Atlas and Health Insight Engine after the bulletin pipeline is trusted."),
            ("PDF distribution", "Build light template export in hardening phase. The sprint proves data correctness first."),
        ],
        title="Build vs. buy decisions",
    )

    add_section_heading(doc, "Section 2: Working Prototype Code", 2)
    add_callout(
        doc,
        "Prototype delivered",
        "The submitted Django project reads the provided healthcare source files, joins them by facility, calculates bulletin "
        "metrics, supports monthly/quarterly/annual filters, and serves both an HTML dashboard and JSON API.",
        fill=LIGHT_TEAL,
    )
    add_metrics_strip(
        doc,
        [
            ("Tests", "9 passing", "Django service/API coverage"),
            ("Annual deliveries", f"{annual_report.maternal_summary['deliveries']:,}", "2024 summed from source months"),
            ("Q4 completeness", f"{quarterly_report.rollup_summary['selected_reporting_completeness_pct']}%", "Partial quarter flagged"),
            ("Advanced metrics", f"{len(annual_report.advanced_aggregations)}", "Rollups, coverage, stability, deltas"),
        ],
    )
    add_bullets(
        doc,
        [
            "Reads the five provided source files: clinical_neonatal, facilities, governance, healthcare_workers, and operations.",
            "Automatically classifies source files, joins facility-level dimensions, and tolerates additional supplemental facility files.",
            "Calculates top facilities, maternal/neonatal indicators, facility performance scores, provincial overview, trend analysis, and risk watchlist.",
            "Adds advanced filtering from lowest report unit upward: monthly source rows, quarterly rollups, and annual rollups.",
            "Recomputes rates from summed counts: annual neonatal mortality uses total neonatal deaths divided by total live births, not averaged quarterly rates.",
            "Exposes outputs at /, /api/bulletin/, and /export/.",
        ],
    )
    add_key_value_table(
        doc,
        [
            ("Run tests", "python manage.py test"),
            ("Run server", "python manage.py runserver"),
            ("Dashboard", "http://127.0.0.1:8000/?granularity=Q&period=2024Q4"),
            ("Annual rollup demo", "http://127.0.0.1:8000/?granularity=A&period=2024"),
            ("JSON API", "http://127.0.0.1:8000/api/bulletin/?granularity=A&period=2024"),
            ("Core code", "bulletin/services.py, bulletin/views.py, bulletin/templates/bulletin/dashboard.html"),
        ],
        title="How to review the working code",
    )

    add_section_heading(doc, "Prototype output evidence", 3)
    add_image_if_exists(
        doc,
        SCREENSHOT_DIR / "dashboard_hero.png",
        "Figure 1 - Executive dashboard, reporting filters, KPI strip, and strategic summary.",
    )
    add_image_if_exists(
        doc,
        SCREENSHOT_DIR / "dashboard_advanced_rollups.png",
        "Figure 2 - Advanced aggregation intelligence and annual rollup explorer.",
    )
    add_image_if_exists(
        doc,
        SCREENSHOT_DIR / "dashboard_tables.png",
        "Figure 3 - Facility tables, risk watchlist, provincial performance, and score leaders.",
    )

    add_section_heading(doc, "Section 3: Implementation Notes", 2)
    add_key_value_table(
        doc,
        [
            ("Shortcuts I took", "Used Google Drive/CSV/XLSX files instead of live DHIS2 API credentials; kept scores in code rather than admin-configurable tables; used JSON export before official PDF automation."),
            ("What works in Week 3", "Ingestion, joins, monthly/quarterly/annual filters, KPIs, rollups, risk alerts, dashboard, JSON API, export route, screenshots, and unit tests."),
            ("What is still not production", "Authentication, scheduled orchestration, signed-off indicator governance, Superset auto-registration, observability, and official MoH PDF distribution."),
            ("What I learned", "The hardest part is not rendering a dashboard; it is making the metric lineage trustworthy enough that MoH leaders can act on it."),
        ],
    )

    add_section_heading(doc, "Deliverable 3: Production Hardening", 1)
    hardening = [
        ("1. DHIS2 integration and orchestration", "CSV exports can recreate manual bottlenecks and stale data.", "Use a scheduled DHIS2 API connector through HealthOS, idempotent sync jobs, retries, freshness badges, and alerting when reporting is late."),
        ("2. Metric governance and data quality", "Hard-coded definitions can drift from MoH-approved indicators.", "Create a metric-definition registry with source fields, formulas, owners, validation rules, and sign-off history."),
        ("3. Security and access control", "Health data must not be exposed uniformly to every user.", "Add MoH SSO, role-based access, district-level row filters, audit logs, HTTPS, and secret management."),
        ("4. Reliability and observability", "Pipeline failures must not silently corrupt a national bulletin.", "Add job monitoring, structured logs, dead-letter files, reconciliation reports, backups, and an L1/L2 incident runbook."),
        ("5. Official outputs and adoption", "A working dashboard is not enough if MoH still needs the official bulletin format.", "Generate branded PDF/Word bulletins, publish a versioned archive, email approved outputs, and connect Superset for analyst drill-down."),
    ]
    for title, concern, answer in hardening:
        add_section_heading(doc, title, 2)
        add_key_value_table(doc, [("Concern", concern), ("How I would address it", answer)])

    add_section_heading(doc, "Deliverable 4: Handover and Deployment", 1)
    add_key_value_table(
        doc,
        [
            ("Data pipeline runbook", "How ingestion runs, where credentials live, how to retry failed jobs, and who owns each escalation step."),
            ("Metric definitions and lineage", "Every bulletin metric mapped to source data, formula, aggregation rule, owner, and known limitation."),
            ("Deployment and operations guide", "Environment setup, backups, monitoring, release process, access control, and manual recovery steps."),
        ],
        title="Three documents I would leave with MoH IT",
    )
    add_key_value_table(
        doc,
        [
            ("MoH Planning analysts", "Run the bulletin, validate numbers, interpret trends, export official outputs, and handle partial periods."),
            ("DHIS2 administrators", "Maintain API tokens, validate data elements, troubleshoot sync errors, and review completeness."),
            ("District Health Officers", "Use risk watchlists, provincial/facility comparisons, and trend signals to plan follow-up."),
            ("MoH IT", "Deploy, monitor, back up, patch, and escalate incidents."),
            ("Sand country team", "Translate adoption blockers into reusable product patterns for HealthOS, Atlas, and Insight Engine."),
        ],
        title="Training plan",
    )
    add_section_heading(doc, "Exit criteria", 2)
    add_bullets(
        doc,
        [
            "MoH staff generate two consecutive bulletins without Sand driving the process.",
            "Metric definitions, runbooks, and deployment guide are signed off by Planning and IT leads.",
            "Automated output is accepted as the source for the official bulletin workflow.",
            "District users can identify and act on high-risk facilities using the dashboard without additional explanation.",
            "Sand transitions from daily implementation support to L2 escalation and product-pattern reuse.",
        ],
    )

    doc.add_page_break()
    add_section_heading(doc, "Appendix A: Evidence from the Current Build", 1)
    add_key_value_table(
        doc,
        [
            ("Latest monthly period", f"{monthly_report.current_period_label}: {monthly_report.maternal_summary['deliveries']:,} deliveries."),
            ("Quarterly period", f"{quarterly_report.current_period_label}: {quarterly_report.maternal_summary['deliveries']:,} deliveries; previous comparison {quarterly_report.previous_period_label}."),
            ("Annual period", f"{annual_report.current_period_label}: {annual_report.maternal_summary['deliveries']:,} deliveries and {annual_report.maternal_summary['neonatal_mortality_rate']}/1k NMR."),
            ("Annual reconciliation", f"Quarter-summed deliveries = {annual_report.rollup_summary['quarter_sum_deliveries']:,}; annual maternal total = {annual_report.maternal_summary['deliveries']:,}."),
            ("Top facility in Q4", f"{quarterly_report.top_facilities[0]['facility_name']} with {quarterly_report.top_facilities[0]['deliveries']:,} deliveries."),
            ("Risk watchlist size", f"{len(quarterly_report.risk_alerts)} facilities flagged for overlapping burden and fragility."),
        ],
    )
    add_section_heading(doc, "Appendix B: Final Submission Package", 1)
    add_bullets(
        doc,
        [
            "Scoping and design document: this Word file.",
            "Working code: Django project in sand-fde-assignment/.",
            "README: setup, test, run, and generation instructions.",
            "Sample data: provided source files under data/google_drive_source/.",
            "Output demonstration: screenshots in submission_assets/ and live dashboard/API routes.",
        ],
    )
    doc.add_paragraph(
        "I wrote this submission to be defendable in the live interview: the tradeoffs are explicit, the code runs, "
        "the metrics are traceable, and the production path ties back to Sand's existing healthcare products."
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT_MAIN)
    shutil.copyfile(OUTPUT_MAIN, OUTPUT_COMPAT)
    try:
        shutil.copyfile(OUTPUT_MAIN, OUTPUT_DOWNLOADS)
    except PermissionError:
        pass
    print(f"Created: {OUTPUT_MAIN}")
    print(f"Updated: {OUTPUT_COMPAT}")
    if OUTPUT_DOWNLOADS.exists():
        print(f"Copied: {OUTPUT_DOWNLOADS}")


if __name__ == "__main__":
    main()
